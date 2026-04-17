# accounts/views.py
from rest_framework import status, permissions, serializers, generics, viewsets
from django.db import models, transaction
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from django.contrib.auth import authenticate
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework.permissions import IsAuthenticated

# Models
from Staff_profile.models import StaffProfile 
from accounts.models import (
    User, Student, Document, TimeTable, Letter, 
    Request, RequestHistory, Notice, NoticeAcknowledgement, NoticeComment, Department, ClassAdvisor, AccountRequest, UserCreationRequest
)

from accounts.serializers import (
    UserSerializer, StudentSerializer, CustomTokenObtainPairSerializer,
    DocumentSerializer, TimeTableSerializer, LetterSerializer,
    RequestSerializer, RequestActionSerializer, AdminActionSerializer, PrincipalActionSerializer,
    NoticeSerializer, NoticeAcknowledgementSerializer, NoticeCommentSerializer,
    UserCreateSerializer, BulkUploadSerializer, DepartmentSerializer, ClassAdvisorSerializer, AccountRequestSerializer, UserCreationRequestSerializer
)
from accounts.views_notifications import create_notification

# Permissions
from accounts.permissions import (
    IsAdmin, IsStaff, IsStudent, IsStaffOrAdmin,
    IsSuperAdmin, IsPrincipal, IsDepartmentAdmin, IsDepartmentStaff
)

class TimeTableListCreateView(generics.ListCreateAPIView):
    queryset = TimeTable.objects.all()
    serializer_class = TimeTableSerializer
    permission_classes = [IsStaffOrAdmin]  # GET allowed for authenticated, POST for staff/admin

    def get_queryset(self):
        # Optionally filter by query params: class_name, section, year
        qs = super().get_queryset()
        class_name = self.request.query_params.get("class_name")
        section = self.request.query_params.get("section")
        year = self.request.query_params.get("year")
        if class_name:
            qs = qs.filter(class_name__iexact=class_name)
        if section:
            qs = qs.filter(section__iexact=section)
        if year:
            qs = qs.filter(year__iexact=year)
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


class TimeTableRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = TimeTable.objects.all()
    serializer_class = TimeTableSerializer
    permission_classes = [IsStaffOrAdmin]  # GET OK for authenticated, PUT/DELETE restricted to staff/admin

    def patch(self, request, *args, **kwargs):
        # allow partial update (e.g. modify grid)
        return self.partial_update(request, *args, **kwargs)
    
class StudentTimeTableListView(generics.ListAPIView):
    queryset = TimeTable.objects.all()
    serializer_class = TimeTableSerializer
    permission_classes = [IsAuthenticated]  # any logged-in student

    def get_queryset(self):
        """
        Students can filter timetables by class_name, section, year.
        Example: /student/timetables/?class_name=IV CSE&section=A&year=2025
        """
        user = self.request.user
        # Only show published timetables
        qs = TimeTable.objects.filter(status='published')
        
        # Filter by student's department if they have one assigned
        if user.department:
            qs = qs.filter(department=user.department)
        
        class_name = self.request.query_params.get("class_name")
        section = self.request.query_params.get("section")
        year = self.request.query_params.get("year")
        
        if class_name:
            qs = qs.filter(class_name__iexact=class_name)
        if section:
            qs = qs.filter(section__iexact=section)
        if year:
            qs = qs.filter(year__iexact=year)
        
        return qs


class StudentTimeTableDetailView(generics.RetrieveAPIView):
    queryset = TimeTable.objects.all()
    serializer_class = TimeTableSerializer
    permission_classes = [IsAuthenticated]  # any logged-in student

# List all timetables created by staff
class TimeTableListView(generics.ListAPIView):
    serializer_class = TimeTableSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return TimeTable.objects.filter(created_by=self.request.user).order_by("-created_at")

class TimeTablePublishView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, pk):
        try:
            timetable = TimeTable.objects.get(pk=pk, created_by=request.user)
            # Allow resubmit if rejected or draft
            if timetable.status not in ['draft', 'rejected', 'pending_verification']:
                 # If already published, maybe allow update? For now restrict.
                 pass

            timetable.status = 'pending_verification'
            timetable.save()
            return Response({"message": "Timetable sent for verification", "status": timetable.status})
        except TimeTable.DoesNotExist:
            return Response({"error": "Timetable not found"}, status=status.HTTP_404_NOT_FOUND)

class TimeTableVerifyView(APIView):
    permission_classes = [IsAuthenticated, IsDepartmentAdmin]

    def post(self, request, pk):
        try:
            # Dept Admin can only verify timetables in their department
            timetable = TimeTable.objects.get(pk=pk, department=request.user.department)
            action = request.data.get("action") # 'approve' or 'reject'
            
            if action == "approve":
                timetable.status = "published"
                timetable.verified_by = request.user
            elif action == "reject":
                timetable.status = "rejected"
                timetable.verified_by = request.user
            else:
                return Response({"error": "Invalid action. Use 'approve' or 'reject'"}, status=status.HTTP_400_BAD_REQUEST)
            
            timetable.save()
            return Response({"message": f"Timetable {timetable.status}", "status": timetable.status})
        except TimeTable.DoesNotExist:
            return Response({"error": "Timetable not found or not in your department"}, status=status.HTTP_404_NOT_FOUND)

class DeptPendingTimeTableView(generics.ListAPIView):
    serializer_class = TimeTableSerializer
    permission_classes = [IsAuthenticated, IsDepartmentAdmin]

    def get_queryset(self):
        if not self.request.user.department:
            return TimeTable.objects.none()
        return TimeTable.objects.filter(
            department=self.request.user.department, 
            status='pending_verification'
        ).order_by("-updated_at")


# Delete a timetable
class TimeTableDeleteView(generics.DestroyAPIView):
    queryset = TimeTable.objects.all()
    serializer_class = TimeTableSerializer
    permission_classes = [IsAuthenticated]

# -------------------------------
# Login View
# -------------------------------
@api_view(["POST"])
def login_view(request):
    username = request.data.get("username")
    password = request.data.get("password")

    user = authenticate(username=username, password=password)

    if user is not None:
        refresh = RefreshToken.for_user(user)
        return Response({
            "access": str(refresh.access_token),
            "refresh": str(refresh),
            "role": user.role,  # send user role too
        }, status=status.HTTP_200_OK)
    return Response({"detail": "Invalid credentials"}, status=status.HTTP_401_UNAUTHORIZED)




# -------------------------------
# Who am I
# -------------------------------
@api_view(["GET"])
@permission_classes([permissions.IsAuthenticated])
def whoami(request):
    return Response({
        "username": request.user.username,
        "role": request.user.role
    })


# -------------------------------
# Dashboards
# -------------------------------
@api_view(["GET"])
@permission_classes([IsAdmin])
def admin_dashboard(request):
    return Response({"message": f"Welcome Admin {request.user.username}"})

@api_view(["GET"])
@permission_classes([IsStaff])
def staff_dashboard(request):
    return Response({"message": f"Welcome Staff {request.user.username}"})

@api_view(["GET"])
@permission_classes([IsStudent])
def student_dashboard(request):
    user = request.user
    data = {
        "message": f"Welcome Student {user.username}",
        "username": user.username,
        "first_name": user.first_name,
        "last_name": user.last_name,
        "role": user.role,
        "department": user.department.name if user.department else None,
    }
    
    # Try to get associated student profile
    if hasattr(user, 'student_account'):
        student_profile = user.student_account
        data["year"] = student_profile.year
        data["course"] = student_profile.course
        data["roll_no"] = student_profile.roll_no
    else:
        data["year"] = None
        data["course"] = None
        data["roll_no"] = None

    return Response(data)

# -------------------------------
# JWT Token View
# -------------------------------
class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer

# Staff upload (no extra restriction, just logged in)
class DocumentUploadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        serializer = DocumentSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(uploaded_by=request.user)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class DocumentListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        documents = Document.objects.all().order_by("-id")  # latest first
        serializer = DocumentSerializer(documents, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
class DocumentDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, pk):
        try:
            doc = Document.objects.get(pk=pk)
            doc.delete()
            return Response({"message": "Document deleted successfully"}, status=status.HTTP_200_OK)
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)




# 1. Create + list (own letters + shared letters)
class LetterListCreateView(generics.ListCreateAPIView):
    serializer_class = LetterSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return Letter.objects.filter(
            models.Q(owner=user) | models.Q(is_shared=True)
        ).distinct()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)


# 2. Update + delete only own letters
class LetterDetailView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = LetterSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Letter.objects.filter(owner=self.request.user)


class CreateRequestView(generics.CreateAPIView):
    queryset = Request.objects.all()
    serializer_class = RequestSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        letter_id = self.request.data.get("letter")
        staff_id = self.request.data.get("staff")
        
        try:
           letter = Letter.objects.get(id=letter_id)
           staff_profile = StaffProfile.objects.get(user__id=staff_id)
        except Letter.DoesNotExist:
           raise serializers.ValidationError({"letter": "Letter not found"})
        except StaffProfile.DoesNotExist:
           raise serializers.ValidationError({"staff": "Staff not found"})
        
        staff_profile = StaffProfile.objects.get(user__id=staff_id)
        req = serializer.save(student=self.request.user, letter=letter, staff=staff_profile)

        RequestHistory.objects.create(
        request=req,
        user=self.request.user,
        action="submitted",
        status="pending"
        )

        # Notify Staff
        create_notification(
            recipient=staff_profile.user,
            title="New Student Request",
            message=f"{self.request.user.username} submitted a new request: {letter.title}",
            target_url="/staff/requests"
        )


# 6. Student sees all requests + status
class StudentRequestsListView(generics.ListAPIView):
    serializer_class = RequestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Request.objects.filter(student=self.request.user)
 


# 2. Staff sees all requests assigned to them
class StaffRequestsListView(generics.ListAPIView):
    serializer_class = RequestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        try:
            staff_profile = StaffProfile.objects.get(user=self.request.user)
        except StaffProfile.DoesNotExist:
            return Request.objects.none()
        
        return Request.objects.filter(staff=staff_profile)


# 3. Staff approves/rejects a request
class StaffActionView(generics.UpdateAPIView):
    queryset = Request.objects.all()
    serializer_class = RequestActionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, *args, **kwargs):
        req = self.get_object()
        serializer = self.get_serializer(req, data=request.data, partial=True)
        if serializer.is_valid():
            req = serializer.save()
            
            if req.staff_status == 'rejected':
                req.rejection_reason = req.staff_comment
                req.save()

            # Record history
            RequestHistory.objects.create(
                request=req,
                user=request.user,
                action=f"Staff {req.staff_status}",
                status=req.staff_status,
            )

            # Notify Student
            create_notification(
                recipient=req.student,
                title=f"Request {req.staff_status.capitalize()}",
                message=f"Your request '{req.letter.title}' was {req.staff_status} by your class advisor.",
                target_url="/student/requests"
            )

            return Response(RequestSerializer(req).data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# 1. Admin sees all staff-approved requests (pending admin action)
class AdminRequestsListView(generics.ListAPIView):
    serializer_class = RequestSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        # Only show requests where staff has approved (Both pending and processed by Admin)
        qs = Request.objects.filter(staff_status="approved").order_by(
            models.Case(models.When(admin_status="pending", then=0), default=1),
            '-created_at'
        )
        user = self.request.user
        
        # If Dept Admin, only show requests from their department users
        if user.is_authenticated and getattr(user, 'is_dept_admin', False):
             if user.department:
                 return qs.filter(student__department=user.department)
             return qs.none() 
             
        # Super Admin sees all? Or maybe restricted? Assuming global view for super admin.
        if user.is_super_admin or user.is_superuser:
            return qs

        return qs.none() # Fallback

# 2. Admin approves/rejects a request
class AdminActionView(generics.UpdateAPIView):
    queryset = Request.objects.all()
    serializer_class = AdminActionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, *args, **kwargs):
        req = self.get_object()
        
        new_status = request.data.get("admin_status")
        comment = request.data.get("admin_comment", "")
        # Frontend should send this flag if Dept Admin chose "Send to Principal"
        forward_to_principal = request.data.get("forward", False) 

        if new_status == 'rejected':
             req.admin_status = 'rejected'
             req.admin_comment = comment
             req.rejection_reason = comment
             req.save()
        
        elif new_status == 'approved':
             req.admin_status = 'approved'
             req.admin_comment = comment
             
             if forward_to_principal:
                 req.principal_status = 'pending'
             else:
                 # Final approval -> Mark principal status as approved too (auto-pass)
                 req.principal_status = 'approved'
             
             req.save()
        else:
             return Response({"error": "Invalid status"}, status=status.HTTP_400_BAD_REQUEST)

        RequestHistory.objects.create(
            request=req,
            user=request.user,
            action=f"Admin {req.admin_status}",
            status=req.admin_status,
        )

        create_notification(
            recipient=req.student,
            title=f"HOD Review: {req.admin_status.capitalize()}",
            message=f"Your request '{req.letter.title}' was {req.admin_status} by the HOD.",
            target_url="/student/requests"
        )

        return Response(RequestSerializer(req).data)
 



# 3. Principal see only requests forwarded by Admin (admin="approved" AND principal="pending")
class PrincipalRequestsListView(generics.ListAPIView):
    serializer_class = RequestSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        # User wants "only show dept admin sended to principal letters"
        # This implies items waiting for Principal's action.
        return Request.objects.filter(admin_status="approved", principal_status="pending").order_by('-created_at')

class PrincipalActionView(generics.UpdateAPIView):
    queryset = Request.objects.all()
    serializer_class = PrincipalActionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, *args, **kwargs):
        req = self.get_object()
        new_status = request.data.get("principal_status")
        comment = request.data.get("principal_comment", "")

        if new_status in ['approved', 'rejected']:
             req.principal_status = new_status
             req.principal_comment = comment
             if new_status == 'rejected':
                 req.rejection_reason = comment
             req.save()
             
             RequestHistory.objects.create(
                request=req,
                user=request.user,
                action=f"Principal {new_status}",
                status=new_status,
            )

             create_notification(
                recipient=req.student,
                title=f"Principal Review: {new_status.capitalize()}",
                message=f"Your request '{req.letter.title}' was {new_status} by the Principal.",
                target_url="/student/requests"
             )

             return Response(RequestSerializer(req).data)
        
        return Response({"error": "Invalid Status"}, status=status.HTTP_400_BAD_REQUEST)


# ✅ List notices for the specific audience
class NoticeListView(generics.ListAPIView):
    serializer_class = NoticeSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        
        # Base filter: Notices for the user's role
        role_filter = models.Q()
        if user.is_dept_admin:
            role_filter = models.Q(target_dept_admin=True)
        elif user.is_dept_staff:
            role_filter = models.Q(target_staff=True)
        elif user.is_student():
            role_filter = models.Q(target_student=True)
        elif user.is_principal or user.is_super_admin:
            # Principal/SuperAdmin can see everything or we can limit to what they are targeted by
            # Usually they see everything to monitor.
            return Notice.objects.all()

        # Combine with department logic
        # 1. Global (no department) OR 2. Specific department matches user
        dept_filter = models.Q(target_department__isnull=True)
        if user.department:
            dept_filter |= models.Q(target_department=user.department)

        # Also allow user to see notices they authored themselves
        return Notice.objects.filter(
            (role_filter & dept_filter) | models.Q(author=user)
        ).distinct().order_by("-created_at")

# ✅ Create notice with role logic
class NoticeCreateView(generics.CreateAPIView):
    queryset = Notice.objects.all()
    serializer_class = NoticeSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        user = self.request.user
        
        # Dept Admin and Staff can only post to their department
        if user.is_dept_admin or user.is_dept_staff:
            notice = serializer.save(author=user, target_department=user.department)
        else:
            # Principal/SuperAdmin can post globally or to specific dept (from data)
            notice = serializer.save(author=user)
            
        # Notify Targeted Users
        self.notify_targeted_users(notice)

    def notify_targeted_users(self, notice):
        # 1. Base user set
        users = User.objects.all()

        # 2. Filter by department if specified
        if notice.target_department:
            users = users.filter(department=notice.target_department)
            
        # 3. Filter by roles
        role_q = models.Q()
        if notice.target_student:
            role_q |= models.Q(role=User.Roles.DEPT_STUDENT)
        if notice.target_staff:
            role_q |= models.Q(role=User.Roles.DEPT_STAFF)
        if notice.target_dept_admin:
            role_q |= models.Q(role=User.Roles.DEPT_ADMIN)
            
        if role_q:
            users = users.filter(role_q)
        else:
            # If no targets selected, don't notify anyone
            return

        # 4. Limit to users with Push Devices for performance
        # Actually create_notification handles push devices, but we still need target user objects
        # To avoid massive loops, we only create notifications for users who actually exist in the target.
        
        # Bulk notify if needed, but for now we'll iterate
        # Optimization: Only notify users who have logged in recently or have devices?
        # For now, standard iteration.
        target_users = users.exclude(id=notice.author.id) if notice.author else users
        
        for recipient in target_users:
            create_notification(
                recipient=recipient,
                title=f"New Notice: {notice.title}",
                message=notice.content[:100] if notice.content else "Open specifically to read more.",
                target_url="/student/notice" if recipient.is_student() else ("/staff/notice" if recipient.is_dept_staff else "/dept_admin/notice")
            )

# Helper for permission check
def check_notice_permission(user, author):
    if user == author:
        return True
    
    LEVELS = {
        "SUPER_ADMIN": 3,
        "PRINCIPAL": 3,
        "DEPT_ADMIN": 2,
        "DEPT_STAFF": 1,
        "DEPT_STUDENT": 0
    }
    
    user_level = LEVELS.get(user.role, 0)
    author_level = LEVELS.get(author.role, 0) if author else 0
    
    return user_level > author_level

# ✅ Update Notice (Edit)
class NoticeUpdateView(generics.UpdateAPIView):
    queryset = Notice.objects.all()
    serializer_class = NoticeSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_update(self, serializer):
        notice = self.get_object()
        if not check_notice_permission(self.request.user, notice.author):
             raise permissions.PermissionDenied("You do not have permission to edit this notice.")
        serializer.save()

# ✅ Delete notice
class NoticeDeleteView(generics.DestroyAPIView):
    queryset = Notice.objects.all()
    serializer_class = NoticeSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_destroy(self, instance):
        if not check_notice_permission(self.request.user, instance.author):
             raise permissions.PermissionDenied("You do not have permission to delete this notice.")
        instance.delete()

# ✅ Acknowledge Notice
class NoticeAcknowledgeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, pk):
        try:
            notice = Notice.objects.get(pk=pk)
            ack, created = NoticeAcknowledgement.objects.get_or_create(
                notice=notice, student=request.user
            )
            if created:
                return Response({"message": "Acknowledged successfully"}, status=status.HTTP_201_CREATED)
            return Response({"message": "Already acknowledged"}, status=status.HTTP_200_OK)
        except Notice.DoesNotExist:
            return Response({"error": "Notice not found"}, status=status.HTTP_404_NOT_FOUND)

# ✅ Comment on Notice
class NoticeCommentView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, pk):
        try:
            notice = Notice.objects.get(pk=pk)
            comment_text = request.data.get("comment")
            if not comment_text:
                return Response({"error": "Comment cannot be empty"}, status=status.HTTP_400_BAD_REQUEST)

            comment = NoticeComment.objects.create(
                notice=notice, student=request.user, comment=comment_text
            )
            return Response(NoticeCommentSerializer(comment).data, status=status.HTTP_201_CREATED)
        except Notice.DoesNotExist:
            return Response({"error": "Notice not found"}, status=status.HTTP_404_NOT_FOUND)
        
# Delete comment
class NoticeCommentDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        try:
            comment = NoticeComment.objects.get(pk=pk, student=request.user)
            comment.delete()
            return Response({"message": "Comment deleted"}, status=status.HTTP_204_NO_CONTENT)
        except NoticeComment.DoesNotExist:
            return Response({"error": "Comment not found or not permitted"}, status=status.HTTP_404_NOT_FOUND)

# Edit comment
class NoticeCommentEditView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, pk):
        try:
            comment = NoticeComment.objects.get(pk=pk, student=request.user)
            new_text = request.data.get("comment")
            if not new_text:
                return Response({"error": "Comment cannot be empty"}, status=status.HTTP_400_BAD_REQUEST)
            
            comment.comment = new_text
            comment.save()
            return Response(NoticeCommentSerializer(comment).data, status=status.HTTP_200_OK)
        except NoticeComment.DoesNotExist:
            return Response({"error": "Comment not found or not permitted"}, status=status.HTTP_404_NOT_FOUND)



# --- 1. Create Dept Admin (Super Admin / Principal) ---
class CreateDeptAdminView(generics.CreateAPIView):
    serializer_class = UserCreateSerializer
    permission_classes = [permissions.IsAuthenticated, IsSuperAdmin | IsPrincipal]

    def perform_create(self, serializer):
        # Force role to DEPT_ADMIN
        serializer.save(role=User.Roles.DEPT_ADMIN)

# --- 2. Create Dept Staff (Dept Admin / Super / Principal) ---
class CreateDeptStaffView(generics.CreateAPIView):
    serializer_class = UserCreateSerializer
    permission_classes = [permissions.IsAuthenticated, IsDepartmentAdmin | IsSuperAdmin | IsPrincipal]

    def perform_create(self, serializer):
        user = self.request.user
        role = User.Roles.DEPT_STAFF
        
        # If Dept Admin, enforce their department
        if getattr(user, 'is_dept_admin', False):
             serializer.save(role=role, department=user.department)
        else:
             # Super Admin can set any department (handled by serializer input)
             serializer.save(role=role)

# --- 3. Create Dept Student (Dept Staff / Dept Admin / Super / Principal) ---
class CreateDeptStudentView(generics.CreateAPIView):
    serializer_class = UserCreateSerializer
    permission_classes = [permissions.IsAuthenticated, IsDepartmentStaff | IsDepartmentAdmin | IsSuperAdmin | IsPrincipal]

    def perform_create(self, serializer):
        user = self.request.user
        role = User.Roles.DEPT_STUDENT
        
        # If Dept Staff or Dept Admin, enforce their department
        if getattr(user, 'is_dept_admin', False) or getattr(user, 'is_dept_staff', False):
             serializer.save(role=role, department=user.department)
        else:
             serializer.save(role=role)

# --- Department Management ---
class DepartmentListCreateView(generics.ListCreateAPIView):
    queryset = Department.objects.all()
    serializer_class = DepartmentSerializer
    permission_classes = [permissions.IsAuthenticated, IsSuperAdmin | IsPrincipal]

class DepartmentDetailView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Department.objects.all()
    serializer_class = DepartmentSerializer
    permission_classes = [permissions.IsAuthenticated, IsSuperAdmin | IsPrincipal | IsDepartmentAdmin]


# --- 4. Bulk Upload View ---
try:
    import pandas as pd
except ImportError:
    pd = None

class BulkUploadUsersView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    
    def post(self, request, *args, **kwargs):
        if pd is None:
             return Response({"error": "pandas library not installed on server"}, status=status.HTTP_501_NOT_IMPLEMENTED)

        serializer = BulkUploadSerializer(data=request.data)
        if not serializer.is_valid():
             return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        file = serializer.validated_data['file']
        target_role = serializer.validated_data.get('role')
        is_preview = request.data.get('preview') == 'true'
        
        # Permission Check
        user = request.user
        
        # 1. HANDLE PREVIEW (Bypass verification checks, just validate file content)
        # Allows Dept Admin/Staff to see if their file is valid before submitting request.
        dept_override = None
        if hasattr(user, 'department') and user.department:
            dept_override = user.department

        if is_preview:
             return self.process_bulk_upload_file(file, request.data.get('role'), True, user, department_override=dept_override)

        # 2. HANDLE ACTUAL UPLOAD
        can_skip_verification = user.is_super_admin or user.is_principal
        
        # If user is NOT super admin/principal, create a request for verification
        if not can_skip_verification:
             # Ensure file is provided
             if not file:
                 return Response({"error": "File is required"}, status=status.HTTP_400_BAD_REQUEST)
             
             # Create Request
             req = UserCreationRequest.objects.create(
                 uploaded_by=user,
                 department=user.department,
                 file=file,
                 status="pending"
             )
             return Response({
                 "message": "Bulk upload request submitted for verification.", 
                 "request_id": req.id
             }, status=status.HTTP_201_CREATED)

        # 3. DIRECT UPLOAD (Super Admin/Principal)
        # (We reuse the core logic but refactored to be callable)
        return self.process_bulk_upload_file(file, request.data.get('role'), False, user)

    # Refactored Core Helper
    def process_bulk_upload_file(self, file, target_role, is_preview, user, department_override=None):
        try:
            # Attempt to read as Excel first
            try:
                df = pd.read_excel(file)
            except Exception as e:
                # Fallback: Try reading as CSV
                try:
                    file.seek(0)
                    df = pd.read_csv(file)
                except Exception as csv_e:
                     # If both fail, return original error
                     raise ValueError(f"Could not read file as Excel or CSV. Details: {str(e)}")

            # Sanitize headers
            df.columns = [str(c).strip().lower() for c in df.columns]
            
            # Flexible Column Mapping
            username_aliases = ['user name', 'reg no', 'register number', 'roll no', 'staff id', 'id', 'student name', 'name']
            if 'username' not in df.columns:
                for alias in username_aliases:
                    if alias in df.columns.tolist(): # Use tolist for safer check
                        df.rename(columns={alias: 'username'}, inplace=True)
                        break
            
            # Additional Mapping for Email if missing (maybe create dummy if allowed? No, safer to require it or map from other field) 
            # (Assuming email is critical_cols)

            if 'first_name' not in df.columns:
                 name_aliases = ['student name', 'name', 'full name', 'staff name']
                 for alias in name_aliases:
                     if alias in df.columns.tolist() and alias != 'username':
                         df.rename(columns={alias: 'first_name'}, inplace=True)
                         break

            df = df.loc[:, ~df.columns.str.contains('^unnamed')]
            critical_cols = [c for c in ['username'] if c in df.columns] # Check at least username exists
            if critical_cols: df.dropna(subset=critical_cols, how='all', inplace=True)
            
            if 'username' not in df.columns:
                 return Response({"error": f"Missing required column: 'username' (or alias like {username_aliases}). Found: {list(df.columns)}"}, status=status.HTTP_400_BAD_REQUEST)
                
        except Exception as e:
            return Response({"error": f"Failed to validate file structure: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)
            
        created_count = 0
        preview_rows = []
        errors = []
        
        # Determine Permissions for VALIDATION context
        # (This executes as Super Admin basically if called from approval, 
        # or as the calling user if skipped verification. 
        # The 'user' passed here is the ACTOR. If verifying, ACTOR is admin.)
        # Logic: If Verifying, we trust the ACTOR (Admin). The Creator was the requester.
        # But we must ensure data is valid.
        
        can_create_admin = user.is_super_admin or user.is_principal
        can_create_staff = can_create_admin or user.is_dept_admin or getattr(user, 'is_dept_staff', False) # Allow staff to create students? The logic allows staff now.
        can_create_student = True # Everyone allowed here can create students

        try:
            with transaction.atomic():
                for index, row in df.iterrows():
                    try:
                        # Use safe get with sanitized keys
                        username = row.get('username')
                        email = row.get('email')
                        password = str(row.get('password', 'Default@123')) 
                        
                        if pd.isna(username) or pd.isna(email):
                            errors.append(f"Row {index}: Missing username or email")
                            continue
                            
                        # Sanitize Username (Handle float/int from Excel)
                        try:
                            if isinstance(username, float) and username.is_integer():
                                username = int(username)
                            username = str(username).strip()
                            # Double check if string still has .0 (e.g. from string cell "123.0")
                            if username.endswith('.0'): 
                                username = username[:-2]
                        except:
                            username = str(username).strip()
                        email = str(email).strip()
                        
                        first_name = row.get('first_name')
                        last_name = row.get('last_name')
                        first_name = str(first_name).strip() if not pd.isna(first_name) else ""
                        last_name = str(last_name).strip() if not pd.isna(last_name) else ""

                        # Determine Role
                        row_role = target_role or row.get('role')
                        if not row_role or pd.isna(row_role):
                            # Default to STUDENT if not specified
                            row_role = User.Roles.DEPT_STUDENT
                        
                        row_raw_role = str(row_role).strip().upper()
                        row_role = row_raw_role

                        role_map = {
                            'STUDENT': User.Roles.DEPT_STUDENT, 'STAFF': User.Roles.DEPT_STAFF,
                            'FACULTY': User.Roles.DEPT_STAFF, 'TEACHER': User.Roles.DEPT_STAFF,
                            'ADMIN': User.Roles.DEPT_ADMIN, 'HOD': User.Roles.DEPT_ADMIN,
                            'DEPT_HOD': User.Roles.DEPT_ADMIN, 'DEPT_STUDENT': User.Roles.DEPT_STUDENT,
                            'DEPT_STAFF': User.Roles.DEPT_STAFF, 'DEPT_ADMIN': User.Roles.DEPT_ADMIN
                        }
                        if row_role in role_map: row_role = role_map[row_role]

                        # REJECTION Logic: Staff can ONLY create Students.
                        # If a Dept Admin/Staff tries to create Staff/Admin -> Reject.
                        # BUT this method is run by Super Admin during approval. 
                        # We must respect the REQUESTER's limits? 
                        # Actually, user requirement: "rejected if not specified it automatically consider as student" - Wait, "staff can only create students they specify role staff or admin it will reject"
                        pass

                        if row_role == User.Roles.DEPT_ADMIN and not can_create_admin:
                             errors.append(f"Row {index}: Permission denied to create Dept Admin")
                             continue
                        if row_role == User.Roles.DEPT_STAFF and not can_create_staff:
                             errors.append(f"Row {index}: Permission denied to create Staff")
                             continue
                        
                        # Data Validation for Students
                        if row_role == User.Roles.DEPT_STUDENT:
                            if not first_name:
                                errors.append(f"Row {index}: First Name is required")
                                continue
                             
                        # Determine Department
                        from accounts.models import Department
                        dept_name = row.get('department')
                        resolved_department = None
                        
                        if department_override:
                            resolved_department = department_override
                        elif dept_name and not pd.isna(dept_name):
                             dept_name = str(dept_name).strip()
                             resolved_department = Department.objects.filter(name__iexact=dept_name).first()
                             # (Aliases lookup omitted for brevity, use same map as before if needed)
                             if not resolved_department:
                                 # Fallback to user's department
                                 resolved_department = user.department

                        if not resolved_department:
                             # Fallback
                             resolved_department = user.department

                        final_department = resolved_department
                        
                        if User.objects.filter(email=email).exists():
                             errors.append(f"Row {index+2}: Email '{email}' already exists")
                             continue
        
                        if User.objects.filter(username=username).exists():
                             errors.append(f"Row {index+2}: Register No '{username}' already exists")
                             continue
                        
                        student_year = None
                        if row_role == User.Roles.DEPT_STUDENT:
                            student_year = row.get('year')
                            if pd.isna(student_year): student_year = 1 # Default? Or Error. Let's not error on year, just optional.
                            try: student_year = int(student_year)
                            except: student_year = None

                        if is_preview:
                            # Enhanced Preview Data
                            preview_data = {
                                "username": username, 
                                "email": email, 
                                "role": row_raw_role,
                                "first_name": first_name, 
                                "last_name": last_name,
                                "department": final_department.name if final_department else "N/A"
                            }
                            
                            # Add Student specific preview fields
                            if row_role == User.Roles.DEPT_STUDENT:
                                semester = row.get('semester'); semester = row.get('sem') if pd.isna(semester) else semester
                                year_v = row.get('year')
                                mobile = row.get('mobile_number'); mobile = row.get('mobile number') if pd.isna(mobile) else mobile
                                
                                if not pd.isna(semester): 
                                    try: preview_data['semester'] = str(int(float(semester))) 
                                    except: preview_data['semester'] = str(semester)
                                if not pd.isna(year_v): 
                                    try: preview_data['year'] = str(int(float(year_v)))
                                    except: preview_data['year'] = str(year_v)
                                if not pd.isna(mobile): 
                                    try: preview_data['mobile'] = str(int(float(mobile)))
                                    except: preview_data['mobile'] = str(mobile)

                            preview_rows.append(preview_data)
                            continue

                        new_user = User.objects.create_user(
                            username=username, email=email, password=password, role=row_role,
                            first_name=first_name, last_name=last_name, department=final_department
                        )
                        
                        if row_role == User.Roles.DEPT_STUDENT:
                            # Extract extra fields
                            semester = row.get('semester'); semester = row.get('sem') if pd.isna(semester) else semester
                            mobile = row.get('mobile_number'); mobile = row.get('mobile number') if pd.isna(mobile) else mobile
                            dob_val = row.get('dob'); dob_val = row.get('date of birth') if pd.isna(dob_val) else dob_val
                            gender_val = row.get('gender'); gender_val = row.get('sex') if pd.isna(gender_val) else gender_val
                            blood = row.get('blood_group'); blood = row.get('blood group') if pd.isna(blood) else blood
                            addr = row.get('address')
                            scholar = row.get('scholar_type')

                            if not pd.isna(semester):
                                try: semester = int(semester)
                                except: semester = None
                            
                            Student.objects.create(
                                user=new_user, roll_no=username, year=student_year,
                                course=final_department.name if final_department else "General",
                                semester=semester,
                                mobile_number=str(mobile) if not pd.isna(mobile) else None,
                                dob=dob_val if not pd.isna(dob_val) else None,
                                gender=str(gender_val) if not pd.isna(gender_val) else None,
                                blood_group=str(blood) if not pd.isna(blood) else None,
                                address=str(addr) if not pd.isna(addr) else None,
                                scholar_type=str(scholar) if not pd.isna(scholar) else None
                            )

                        assign_role_permissions(new_user)
                        created_count += 1
                        
                    except Exception as e:
                        import traceback; traceback.print_exc()
                        errors.append(f"Row {index}: Unexpected error: {str(e)}")
                
                if is_preview:
                    return Response({
                        "preview": True, "users": preview_rows,
                        "valid_count": len(preview_rows), "errors": errors
                    }, status=status.HTTP_200_OK)

                if errors:
                    transaction.set_rollback(True)
                    return Response({"message": "Errors found.", "errors": errors}, status=status.HTTP_400_BAD_REQUEST)

        except Exception as e:
             import traceback; traceback.print_exc()
             return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
            
        return Response({"message": f"Successfully created {created_count} users"}, status=status.HTTP_201_CREATED)

class UserCreationRequestPreviewView(BulkUploadUsersView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk, *args, **kwargs):
        # Fetch request
        try:
            req = UserCreationRequest.objects.get(pk=pk)
        except UserCreationRequest.DoesNotExist:
            return Response({"error": "Request not found"}, status=status.HTTP_404_NOT_FOUND)

        # Check permission
        user = request.user
        if not (user.is_super_admin or user.is_principal or (user.is_dept_admin and req.department == user.department)):
             return Response({"error": "Permission denied"}, status=status.HTTP_403_FORBIDDEN)

        try:
             req.file.open('rb')
             # Use req.department as override to simulate the upload context correctly
             return self.process_bulk_upload_file(
                 req.file, 
                 None, 
                 True, 
                 user, 
                 department_override=req.department
             )
        except Exception as e:
             return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

# Admin: List Requests
class UserCreationRequestListView(generics.ListAPIView):
    serializer_class = UserCreationRequestSerializer
    permission_classes = [permissions.IsAuthenticated] # Allow Dept Admin too

    def get_queryset(self):
        user = self.request.user
        
        # Super Admin / Principal: See pending requests, BUT exclude Staff requests (hierarchy: Staff -> Dept Admin -> Principal)
        if user.is_super_admin or user.is_principal:
            return UserCreationRequest.objects.filter(status="pending") \
                .exclude(uploaded_by__role='DEPT_STAFF').order_by('-created_at')
        
        if user.is_dept_admin:
             # Dept Admin sees PENDING requests from Staff in their department
             # They verify Staff uploads.
             return UserCreationRequest.objects.filter(
                 status="pending", 
                 department=user.department
             ).exclude(uploaded_by=user).order_by('-created_at')

        # Staff: See their own upload history (Pending, Rejected, Approved)
        if user.role == "DEPT_STAFF" or user.is_dept_staff:
            return UserCreationRequest.objects.filter(uploaded_by=user).order_by('-created_at')
        
        return UserCreationRequest.objects.none()

# Admin: Act on Request
class UserCreationRequestActionView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, pk):
        user = request.user
        
        try:
            req = UserCreationRequest.objects.get(pk=pk)
        except UserCreationRequest.DoesNotExist:
             return Response({"error": "Request not found"}, status=status.HTTP_404_NOT_FOUND)
        
        # PERMISSION CHECK
        # 1. Super Admin / Principal can act on ANY
        # 2. Dept Admin can act on THEIR DEPT's requests only
        
        is_authorized = False
        if user.is_super_admin or user.is_principal:
            is_authorized = True
        elif user.is_dept_admin and req.department == user.department:
            is_authorized = True
        
        if not is_authorized:
             return Response({"error": "Permission denied"}, status=status.HTTP_403_FORBIDDEN)

        action = request.data.get('action') # 'approve' or 'reject'
        
        if action == 'reject':
            req.status = 'rejected'
            comment = request.data.get('comment', 'No reason provided')
            req.admin_comment = f"Rejected by {user.username} ({user.role}): {comment}"
            req.save()
            
            # Notify Requester
            create_notification(
                recipient=req.uploaded_by,
                title="Bulk Upload Rejected",
                message=f"Your bulk upload request for {req.department.name if req.department else 'N/A'} was rejected.",
                target_url="/dept_admin/requests" if req.uploaded_by.role == User.Roles.DEPT_ADMIN else "/staff/requests"
            )
            
            return Response({"message": "Request rejected"})
        
        elif action == 'approve':
            # Execute Upload
            view = BulkUploadUsersView()
            
            response = view.process_bulk_upload_file(
                req.file, 
                target_role=None, 
                is_preview=False, 
                user=request.user, 
                department_override=req.department
            )
            
            if response.status_code == 201:
                req.status = 'approved'
                req.admin_comment = f"Approved by {user.username} ({user.role})"
                req.save()
                
                # Notify Requester
                create_notification(
                    recipient=req.uploaded_by,
                    title="Bulk Upload Approved",
                    message=f"Your bulk upload request for {req.department.name if req.department else 'N/A'} was approved.",
                    target_url="/dept_admin/requests" if req.uploaded_by.role == User.Roles.DEPT_ADMIN else "/staff/requests"
                )
                
                return response
            else:
                return response 
        
        return Response({"error": "Invalid action"}, status=status.HTTP_400_BAD_REQUEST)

class UserCreationRequestDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        try:
             req = UserCreationRequest.objects.get(pk=pk)
             if req.uploaded_by != request.user:
                 return Response({"error": "Permission denied"}, status=status.HTTP_403_FORBIDDEN)
             req.delete()
             return Response({"message": "Request deleted successfully"}, status=status.HTTP_204_NO_CONTENT)
        except UserCreationRequest.DoesNotExist:
             return Response({"error": "Request not found"}, status=status.HTTP_404_NOT_FOUND)

class UserCreationRequestClearView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request):
        # Delete all requests uploaded by the current user
        count, _ = UserCreationRequest.objects.filter(uploaded_by=request.user).delete()
        return Response({"message": f"Deleted {count} requests"}, status=status.HTTP_200_OK)

def assign_role_permissions(user):
    """
    Helper to assign default permissions based on role.
    Replicates logic from CustomUserAdmin.save_model
    """
    from django.contrib.auth.models import Permission
    from django.contrib.contenttypes.models import ContentType
    from accounts.models import Student, Notice, Document, Letter, Request, Department, User

    if user.role == User.Roles.DEPT_ADMIN:
        # Full Management: User, Student, Notice, Document, Letter, Request
        manage_models = [User, Student, Notice, Document, Letter, Request]
        perms_to_add = []
        
        for model_cls in manage_models:
            ct = ContentType.objects.get_for_model(model_cls)
            p_list = Permission.objects.filter(content_type=ct, codename__in=[
                f'add_{model_cls._meta.model_name}',
                f'change_{model_cls._meta.model_name}',
                f'delete_{model_cls._meta.model_name}',
                f'view_{model_cls._meta.model_name}'
            ])
            perms_to_add.extend(p_list)
            
        # View Department
        dept_ct = ContentType.objects.get_for_model(Department)
        perms_to_add.extend(Permission.objects.filter(content_type=dept_ct, codename='view_department'))
        
        user.user_permissions.add(*perms_to_add)

    elif user.role == User.Roles.DEPT_STAFF:
        # Manage Student
        student_ct = ContentType.objects.get_for_model(Student)
        student_perms = Permission.objects.filter(content_type=student_ct, codename__in=[
            'add_student', 'change_student', 'delete_student', 'view_student'
        ])
        
        # Manage User (Add/Change/View)
        user_ct = ContentType.objects.get_for_model(User)
        user_perms = Permission.objects.filter(content_type=user_ct, codename__in=[
            'add_user', 'change_user', 'view_user' # Removed delete_user for Staff safety usually, but matches admin logic
        ])
        
        user.user_permissions.add(*student_perms)
        user.user_permissions.add(*user_perms)

class ParseTimetableView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file provided"}, status=400)

        filename = file_obj.name.lower()
        
        try:
            data = {}
            if filename.endswith('.xlsx') or filename.endswith('.xls'):
                data = self.parse_excel(file_obj)
            elif filename.endswith('.pdf'):
                data = self.parse_pdf(file_obj)
            elif filename.endswith('.docx') or filename.endswith('.doc'):
                data = self.parse_word(file_obj)
            else:
                return Response({"error": "Unsupported file format"}, status=400)

            return Response(data, status=200)
        except Exception as e:
            print(f"Parsing Error: {e}")
            import traceback
            traceback.print_exc()
            return Response({"error": f"Failed to parse file: {str(e)}"}, status=400)

    def parse_excel(self, file_obj):
        import pandas as pd
        
        # Load all sheets or just the first one
        df = pd.read_excel(file_obj, header=None)
        
        # --- Common Containers ---
        metadata = {}
        courses = []
        grid_data = [['' for _ in range(7)] for _ in range(6)] 
        college_name = ""
        doc_title = ""
        
        # 1. Metadata (Text Scan)
        top_rows = df.head(20).fillna("").astype(str)
        for idx, row in top_rows.iterrows():
            row_text = " ".join(row.values).lower()
            self._extract_metadata_from_text(row_text, metadata)
            if "college" in row_text and not college_name:
                for cell in row.values:
                    if "college" in cell.lower(): college_name = cell.strip().upper(); break
            if "time table" in row_text: doc_title = "TIME TABLE"

        # 2. Courses (Header Search)
        course_start_idx = -1
        course_cols = {}
        for idx, row in df.fillna("").astype(str).iterrows():
            row_vals = [x.lower().strip() for x in row.values]
            if "code" in row_vals and ("title" in row_vals or "name" in row_vals):
                course_start_idx = idx
                for c_i, val in enumerate(row_vals):
                    if "code" in val: course_cols['code'] = c_i
                    if "title" in val or "name" in val: course_cols['name'] = c_i
                    if "faculty" in val or "staff" in val: course_cols['faculty'] = c_i
                    if "acronym" in val or "abbr" in val: course_cols['acronym'] = c_i
                break
        
        if course_start_idx != -1:
            for idx in range(course_start_idx + 1, len(df)):
                row = df.iloc[idx].fillna("").astype(str).values
                code = row[course_cols['code']].strip() if 'code' in course_cols else ""
                if not code or "total" in code.lower(): break
                courses.append({
                    "id": str(len(courses) + 1),
                    "code": code,
                    "name": row[course_cols['name']].strip() if 'name' in course_cols else "",
                    "faculty": row[course_cols['faculty']].strip() if 'faculty' in course_cols else "",
                    "acronym": row[course_cols['acronym']].strip() if 'acronym' in course_cols else code,
                    "periodsPerWeek": "0"
                })

        # 3. Grid (MON/TUE Search)
        grid_start_idx = -1
        day_col_idx = 0
        days_map = ["MON", "TUE", "WED", "THU", "FRI", "SAT"]
        
        for idx, row in df.fillna("").astype(str).iterrows():
            for c_i in range(3): 
                if c_i < len(row.values) and "MON" in row.values[c_i].upper():
                    grid_start_idx = idx; day_col_idx = c_i; break
            if grid_start_idx != -1: break
            
        if grid_start_idx != -1:
            for d_i, day_name in enumerate(days_map):
                curr_row = grid_start_idx + d_i
                if curr_row >= len(df): break
                row_vals = df.iloc[curr_row].fillna("").astype(str).values
                candidates = row_vals[day_col_idx+1:]
                self._fill_grid_row(grid_data, d_i, candidates)

        return { "metadata": metadata, "courses": courses, "gridData": grid_data, "collegeName": college_name, "docTitle": doc_title }

    def parse_pdf(self, file_obj):
        import pdfplumber
        
        metadata = {}
        courses = []
        grid_data = [['' for _ in range(7)] for _ in range(6)]
        college_name = "" 
        doc_title = ""

        with pdfplumber.open(file_obj) as pdf:
            first_page = pdf.pages[0]
            text = first_page.extract_text()
            
            # Metadata
            if text:
                lines = text.split('\n')
                for line in lines:
                    self._extract_metadata_from_text(line.lower(), metadata)
                    if "college" in line.lower() and not college_name: college_name = line.strip().upper()
                    if "time table" in line.lower(): doc_title = "TIME TABLE"

            # Tables
            tables = first_page.extract_tables()
            
            # Heuristic: Find which table is which
            # Grid table usually has "MON", "TUE" in first column
            # Course table usually has "Code", "Title" in headers
            
            for table in tables:
                if not table: continue
                
                # Check Header
                header = [str(x).lower() for x in table[0] if x]
                
                # Is Course Table?
                if any("code" in x for x in header) and any("title" in x or "name" in x for x in header):
                    # Parse Courses
                    # Assume column mapping based on standard order if headers not perfect
                    # Simple map: Code, Title, Faculty
                    c_idx = next((i for i, x in enumerate(header) if "code" in x), 0)
                    t_idx = next((i for i, x in enumerate(header) if "title" in x or "name" in x), 1)
                    f_idx = next((i for i, x in enumerate(header) if "faculty" in x or "staff" in x), 2)
                    
                    for row in table[1:]:
                        if not row or len(row) < 2: continue
                        code = row[c_idx] if c_idx < len(row) and row[c_idx] else ""
                        if not code or "total" in str(code).lower(): continue
                        
                        code = str(code).strip()
                        name = str(row[t_idx]).strip() if t_idx < len(row) and row[t_idx] else ""
                        faculty = str(row[f_idx]).strip() if f_idx < len(row) and row[f_idx] else ""
                        
                        courses.append({
                            "id": str(len(courses)+1),
                            "code": code,
                            "name": name,
                            "faculty": faculty,
                            "acronym": code, # PDF often doesn't have acronym column
                            "periodsPerWeek": "0"
                        })

                # Is Grid Table?
                # Check first Column for DAYS
                first_col = [str(row[0]).upper() for row in table if row and row[0]]
                if "MON" in first_col or "MONDAY" in first_col:
                    days_map = ["MON", "TUE", "WED", "THU", "FRI", "SAT"]
                    d_idx = 0
                    for row in table:
                        if not row: continue
                        first_cell = str(row[0]).upper().strip()
                        if any(x in first_cell for x in ["MON", "TUE", "WED", "THU", "FRI", "SAT"]):
                             # Found a day row
                             # Find index in map
                             matched_day_idx = -1
                             for i, d in enumerate(days_map):
                                 if d in first_cell: matched_day_idx = i; break
                             
                             if matched_day_idx != -1:
                                 # Fill Grid
                                 candidates = row[1:] # Skip day cell
                                 self._fill_grid_row(grid_data, matched_day_idx, candidates)

        return { "metadata": metadata, "courses": courses, "gridData": grid_data, "collegeName": college_name, "docTitle": doc_title }

    def parse_word(self, file_obj):
        import docx
        
        doc = docx.Document(file_obj)
        metadata = {}
        courses = []
        grid_data = [['' for _ in range(7)] for _ in range(6)]
        college_name = ""
        doc_title = ""

        # Metadata from Paragraphs
        for para in doc.paragraphs:
            text = para.text.lower()
            if text.strip():
                self._extract_metadata_from_text(text, metadata)
                if "college" in text and not college_name: college_name = para.text.strip().upper()
                if "time table" in text: doc_title = "TIME TABLE"

        # Tables
        for table in doc.tables:
            # Heuristic checks similar to PDF
            rows = table.rows
            if not rows: continue
            
            # Check Header
            header_cells = [cell.text.lower().strip() for cell in rows[0].cells]
            
            # Course Table?
            if any("code" in x for x in header_cells) and any("title" in x for x in header_cells):
                 c_idx = next((i for i, x in enumerate(header_cells) if "code" in x), 0)
                 t_idx = next((i for i, x in enumerate(header_cells) if "title" in x or "name" in x), 1)
                 f_idx = next((i for i, x in enumerate(header_cells) if "faculty" in x or "staff" in x), 2)

                 for row in rows[1:]:
                     cells = row.cells
                     code = cells[c_idx].text.strip() if c_idx < len(cells) else ""
                     if not code or "total" in code.lower(): continue
                     courses.append({
                         "id": str(len(courses)+1),
                         "code": code,
                         "name": cells[t_idx].text.strip() if t_idx < len(cells) else "",
                         "faculty": cells[f_idx].text.strip() if f_idx < len(cells) else "",
                         "acronym": code,
                         "periodsPerWeek": "0"
                     })
            
            # Grid Table?
            # Check first column of first few rows
            first_col_txt = [row.cells[0].text.upper().strip() for row in rows[:5] if row.cells]
            if any("MON" in x for x in first_col_txt):
                days_map = ["MON", "TUE", "WED", "THU", "FRI", "SAT"]
                for row in rows:
                    if not row.cells: continue
                    first_cell = row.cells[0].text.upper().strip()
                    matched_day_idx = -1
                    for i, d in enumerate(days_map):
                        if d in first_cell: matched_day_idx = i; break
                    
                    if matched_day_idx != -1:
                        candidates = [c.text for c in row.cells[1:]]
                        self._fill_grid_row(grid_data, matched_day_idx, candidates)

        return { "metadata": metadata, "courses": courses, "gridData": grid_data, "collegeName": college_name, "docTitle": doc_title }

    def _extract_metadata_from_text(self, text, metadata):
        import re
        if "academic year" in text:
            match = re.search(r'\d{4}\s*-\s*\d{4}', text)
            if match: metadata['academicYear'] = match.group(0)
        
        if "semester" in text:
             parts = text.split("semester")
             if len(parts) > 1:
                 # Clean up: " : V" -> "V"
                 val = parts[1].strip()
                 if val.startswith(":") or val.startswith("-"): val = val[1:].strip()
                 metadata['semester'] = val.upper()
        
        if "department" in text:
             parts = text.split("department")
             if len(parts) > 1:
                 val = parts[1].replace("of", "").strip()
                 if val.startswith(":") or val.startswith("-"): val = val[1:].strip()
                 metadata['department'] = val.upper()

    def _fill_grid_row(self, grid_data, row_idx, candidates):
        # Helper to fill periods
        p_filled = 0
        for val in candidates:
            # Handle None or non-string
            s_val = str(val).strip()
            if not s_val: continue
            if s_val.upper() in ["BREAK", "LUNCH", "B", "L"]: continue
            
            if p_filled < 7:
                grid_data[row_idx][p_filled] = s_val
                p_filled += 1

# --- Class Advisors ---
class ClassAdvisorListCreateView(generics.ListCreateAPIView):
    serializer_class = ClassAdvisorSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.is_super_admin or user.is_principal:
             return ClassAdvisor.objects.all()
        if getattr(user, 'is_dept_admin', False) or getattr(user, 'is_dept_staff', False):
             if user.department:
                 return ClassAdvisor.objects.filter(department=user.department)
        return ClassAdvisor.objects.none()

    def perform_create(self, serializer):
        user = self.request.user
        from rest_framework.exceptions import ValidationError
        
        # Explicitly check role to ensure department is assigned
        if user.role == User.Roles.DEPT_ADMIN:
             if not user.department:
                 raise ValidationError({"department": "Your account is not linked to any department."})
             serializer.save(department=user.department)
        elif user.role in [User.Roles.SUPER_ADMIN, User.Roles.PRINCIPAL]:
             # Admins might pass department in body, or we let serializer handle it if valid
             # But here we removed department from input fields, so this might fail if they need to set it.
             # For now, focus on Dept Admin success.
             serializer.save()
        else:
             # Fallback
             serializer.save()

class ClassAdvisorRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = ClassAdvisor.objects.all()
    serializer_class = ClassAdvisorSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        user = self.request.user
        if user.is_super_admin or user.is_principal:
             return ClassAdvisor.objects.all()
        if getattr(user, 'is_dept_admin', False):
             if user.department:
                 return ClassAdvisor.objects.filter(department=user.department)
        return ClassAdvisor.objects.none()

class DepartmentStaffListView(generics.ListAPIView):
    serializer_class = UserSerializer 
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        dept_id = self.request.query_params.get('dept_id')

        if user.is_super_admin or user.is_principal:
             qs = User.objects.filter(role=User.Roles.DEPT_STAFF)
             if dept_id:
                  qs = qs.filter(department_id=dept_id)
             return qs
        
        if getattr(user, 'is_dept_admin', False) and user.department:
            return User.objects.filter(department=user.department, role=User.Roles.DEPT_STAFF)
            
        return User.objects.none()

class DepartmentAdminListView(generics.ListAPIView):
    serializer_class = UserSerializer 
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        dept_id = self.request.query_params.get('dept_id')

        if user.is_super_admin or user.is_principal:
             qs = User.objects.filter(role=User.Roles.DEPT_ADMIN)
             if dept_id:
                  qs = qs.filter(department_id=dept_id)
             return qs
            
        return User.objects.none()

class DepartmentStudentListView(generics.ListAPIView):
    serializer_class = UserSerializer 
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        dept_id = self.request.query_params.get('dept_id')

        if user.is_super_admin or user.is_principal:
             qs = User.objects.filter(role=User.Roles.DEPT_STUDENT)
             if dept_id:
                  qs = qs.filter(department_id=dept_id)
             return qs
        
        if getattr(user, 'is_dept_admin', False) and user.department:
            return User.objects.filter(department=user.department, role=User.Roles.DEPT_STUDENT)
            
        return User.objects.none()

class StudentClassAdvisorListView(APIView):
    permission_classes = [permissions.IsAuthenticated, IsStudent]

    def get(self, request):
        user = request.user
        if not user.department:
            return Response({"error": "No department assigned to student"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get year from student account
        try:
             student_profile = user.student_account
             year = student_profile.year
        except Student.DoesNotExist:
             return Response({"error": "Student profile not found"}, status=status.HTTP_404_NOT_FOUND)

        if not year:
             return Response({"error": "Year not assigned to student"}, status=status.HTTP_400_BAD_REQUEST)

        # Find Class Advisor entry
        advisors_list = []
        try:
            class_advisor = ClassAdvisor.objects.get(department=user.department, year=year)
            
            # Serialize specific fields or use simple dict
            if class_advisor.advisor1:
                 advisors_list.append({
                     "id": class_advisor.advisor1.id,
                     "username": class_advisor.advisor1.username,
                     "avatar_url": class_advisor.advisor1.staffprofile.avatar.url if hasattr(class_advisor.advisor1, 'staffprofile') and class_advisor.advisor1.staffprofile.avatar else None
                 })
            
            if class_advisor.advisor2:
                 advisors_list.append({
                     "id": class_advisor.advisor2.id,
                     "username": class_advisor.advisor2.username,
                     "avatar_url": class_advisor.advisor2.staffprofile.avatar.url if hasattr(class_advisor.advisor2, 'staffprofile') and class_advisor.advisor2.staffprofile.avatar else None
                 })

        except ClassAdvisor.DoesNotExist:
            # Fallback: Return all staff in the department if no specific advisor is assigned
            staff_members = User.objects.filter(department=user.department, role=User.Roles.DEPT_STAFF)
            for staff in staff_members:
                 avatar_url = None
                 if hasattr(staff, 'staffprofile') and staff.staffprofile.avatar:
                     avatar_url = staff.staffprofile.avatar.url
                 
                 advisors_list.append({
                     "id": staff.id,
                     "username": staff.username,
                     "avatar_url": avatar_url
                 })
            
        return Response(advisors_list, status=status.HTTP_200_OK)

from accounts.serializers import PrincipalActionSerializer

class PrincipalRequestsListView(generics.ListAPIView):
    serializer_class = RequestSerializer
    permission_classes = [permissions.IsAuthenticated, IsPrincipal]

    def get_queryset(self):
        # Show requests where Dept Admin approved AND Principal status is pending
        return Request.objects.filter(admin_status="approved", principal_status="pending")

class PrincipalActionView(generics.UpdateAPIView):
    queryset = Request.objects.all()
    serializer_class = PrincipalActionSerializer
    permission_classes = [permissions.IsAuthenticated, IsPrincipal]

    def patch(self, request, *args, **kwargs):
        req = self.get_object()
        
        status_val = request.data.get("principal_status")
        comment = request.data.get("principal_comment", "")

        if status_val == "rejected":
             req.principal_status = "rejected"
             req.principal_comment = comment
             req.rejection_reason = comment
        elif status_val == "approved":
             req.principal_status = "approved"
             req.principal_comment = comment
        else:
             return Response({"error": "Invalid status"}, status=status.HTTP_400_BAD_REQUEST)

        req.save()
        
        RequestHistory.objects.create(
            request=req,
            user=request.user,
            action=f"Principal {req.principal_status}",
            status=req.principal_status,
        )
        return Response(RequestSerializer(req).data)

# --- 5. Account Creation Request (Public) ---

@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def public_department_list(request):
    """
    Publicly list departments so users can select one for account request.
    Query Param: category=UG or PG
    """
    category = request.query_params.get('category')
    qs = Department.objects.all()
    if category:
        qs = qs.filter(category=category.upper())
    
    serializer = DepartmentSerializer(qs, many=True)
    return Response(serializer.data)

@api_view(['POST'])
@permission_classes([permissions.AllowAny])
def create_account_request(request):
    """
    Public endpoint to submit an account creation request.
    """
    serializer = AccountRequestSerializer(data=request.data)
    if serializer.is_valid():
        req = serializer.save()
        
        # Notify Dept Admin of this department
        dept_admins = User.objects.filter(department=req.department, role=User.Roles.DEPT_ADMIN)
        for admin in dept_admins:
            create_notification(
                recipient=admin,
                title="New Account Request",
                message=f"{req.full_name} ({req.register_number}) requested an account for {req.department.name}.",
                target_url="/dept_admin/requests"
            )
            
        return Response({"message": "Request submitted successfully. Admin will review it."}, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class DeptAdminAccountRequestListView(generics.ListAPIView):
    serializer_class = AccountRequestSerializer
    permission_classes = [permissions.IsAuthenticated] # Add IsDeptAdmin check ideally

    def get_queryset(self):
        user = self.request.user
        # Strict check: Must be Dept Admin and have a department
        if getattr(user, 'is_dept_admin', False) and user.department:
            return AccountRequest.objects.filter(department=user.department).order_by('-created_at')
        return AccountRequest.objects.none()

class DeptAdminAccountRequestActionView(APIView):
    permission_classes = [permissions.IsAuthenticated] # Add IsDeptAdmin check

    def post(self, request, pk):
        try:
            req = AccountRequest.objects.get(pk=pk, department=request.user.department)
        except AccountRequest.DoesNotExist:
             return Response({"error": "Request not found or not in your department"}, status=status.HTTP_404_NOT_FOUND)

        action = request.data.get("action")
        note = request.data.get("note", "")

        if action == "approve":
             req.status = "approved"
             # Optionally AUTO-CREATE USER? 
             # For now, just mark approved. The Dept Admin might manually create or we can automate.
             # User said: "dept admin get request to create account ... show to an dept admin in requests"
             # It implies they just see the request initially. Automating creation is a bonus step.
        elif action == "reject":
             req.status = "rejected"
        else:
            return Response({"error": "Invalid action"}, status=status.HTTP_400_BAD_REQUEST)
        
        req.resolution_note = note
        req.save()
        return Response(AccountRequestSerializer(req).data)

class ChangePasswordView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        user = request.user
        old_password = request.data.get("old_password")
        new_password = request.data.get("new_password")

        if not old_password or not new_password:
            return Response({"error": "old_password and new_password are required."}, status=status.HTTP_400_BAD_REQUEST)

        if not user.check_password(old_password):
            return Response({"error": "Incorrect old password."}, status=status.HTTP_400_BAD_REQUEST)

        user.set_password(new_password)
        user.save()
        return Response({"message": "Password changed successfully."}, status=status.HTTP_200_OK)

class ActiveSessionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user_agent = request.META.get('HTTP_USER_AGENT', 'Unknown Device')
        ip_address = request.META.get('REMOTE_ADDR', '127.0.0.1')
        
        # Simplified parser for device info
        # Enhanced parser for mobile/app detection
        os = "Unknown"
        device_name = "Desktop Browser"
        
        # Check for common mobile strings
        is_mobile = any(x in user_agent for x in ["Android", "iPhone", "iPad", "Mobile", "Expo", "okhttp", "Darwin"])

        if "Android" in user_agent or "okhttp" in user_agent:
            os = "Android"
            device_name = "Android Mobile"
        elif any(x in user_agent for x in ["iPhone", "iPad", "Darwin"]):
            os = "iOS"
            device_name = "Apple Mobile"
        elif "Windows" in user_agent:
            os = "Windows"
            device_name = "Windows PC"
        elif "Mac" in user_agent:
            os = "macOS"
            device_name = "Mac"
        elif "Linux" in user_agent:
            os = "Linux"
            device_name = "Linux PC"
            
        if is_mobile and device_name == "Desktop Browser":
            device_name = "Mobile App"


        sessions = [
            {
                "id": "current",
                "device_name": device_name,
                "os": os,
                "ip_address": ip_address,
                "last_active": "Just now",
                "is_current": True
            }
        ]
        return Response(sessions, status=status.HTTP_200_OK)
